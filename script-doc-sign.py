import requests
import base64
from docx import Document
from openpyxl import load_workbook
from fpdf import FPDF
import copy
import tempfile
import os

# Função para converter o documento Word em PDF
def convert_to_pdf(word_file):
    pdf = FPDF()
    doc = Document(word_file)
    pdf_content = ""

    for paragraph in doc.paragraphs:
        pdf_content += paragraph.text + "\n"

    return pdf_content

# Função para enviar o PDF para o DocuSign
def send_to_docusign(pdf_content, signer_name, signer_email):
    # Configuração da API do DocuSign
    base_url = 'https://demo.docusign.net/restapi/v2'
    integrator_key = 'YOUR_INTEGRATOR_KEY'
    username = 'YOUR_DOCUSIGN_USERNAME'
    password = 'YOUR_DOCUSIGN_PASSWORD'
    
    # Autenticação na API do DocuSign
    auth_headers = {
        'X-DocuSign-Authentication': f'{{ "Username": "{username}", "Password": "{password}", "IntegratorKey": "{integrator_key}" }}'
    }

    # Enviar o PDF para o DocuSign para a assinatura
    pdf_bytes = pdf_content.encode('utf-8')
    pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')

    request_data = {
        'documents': [{
            'documentBase64': pdf_base64,
            'name': 'Document.pdf',
            'fileExtension': 'pdf',
            'documentId': '1'
        }],
        'recipients': {
            'signers': [{
                'email': signer_email,
                'name': signer_name,
                'recipientId': '1',
                'clientUserId': '1001',
                'tabs': {
                    'signHereTabs': [{
                        'documentId': '1',
                        'pageNumber': '1',
                        'xPosition': '100',
                        'yPosition': '100'
                    }]
                }
            }]
        },
        'status': 'sent'
    }

    response = requests.post(f'{base_url}/envelopes', json=request_data, headers=auth_headers)
    response_data = response.json()
    
    if response.ok:
        envelope_id = response_data['envelopeId']
        return envelope_id
    else:
        error_message = response_data['message']
        print(f'Error: {error_message}')
        return None

if __name__ == '__main__':
    # Carrega a planilha Excel
    workbook = load_workbook('./nomes.xlsx')
    sheet = workbook.active

    # Loop para preencher e enviar os documentos
    for row in sheet.iter_rows(min_row=2, values_only=True):
        nome = row[0]
        cpf = row[1]
        endereco = row[2]
        email = row[3]  # Supondo que o e-mail está na coluna 4 da planilha

        # Carrega o documento Word
        document = Document('./teste.docx')

        # Cria uma cópia do documento para preencher
        filled_document = copy.deepcopy(document)

        # Preenche os campos no documento Word
        for paragraph in filled_document.paragraphs:
            if '<<NOME>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<NOME>>', nome)
            if '<<CPF>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<CPF>>', str(cpf))
            if '<<ENDERECO>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<ENDERECO>>', endereco)

            # Adiciona a assinatura
            if '<<ASSINATURA>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<ASSINATURA>>', '[ASSINATURA]')
        
        # Salva o documento preenchido em um arquivo temporário
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        filled_document.save(temp_file.name)

        # Obter o conteúdo do PDF
        pdf_content = convert_to_pdf(temp_file.name)

        # Enviar o PDF para o DocuSign com o nome e e-mail do signatário
        envelope_id = send_to_docusign(pdf_content, nome, email)

        if envelope_id:
            print(f'Success! Envelope ID: {envelope_id}')
        else:
            print('Error sending the document to DocuSign.')

        # Exclui o arquivo temporário
        os.remove(temp_file.name)
