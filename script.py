from docx import Document
from openpyxl import load_workbook
from fpdf import FPDF
import copy
import tempfile
import os

# Função para converter o documento Word em PDF
def convert_to_pdf(word_file, pdf_file):
    pdf = FPDF()
    doc = Document(word_file)

    for paragraph in doc.paragraphs:
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, txt=paragraph.text, ln=1)

    pdf.output(pdf_file)

# Carrega a planilha Excel
workbook = load_workbook('./nomes.xlsx')
sheet = workbook.active

# Preenche o formulário com os dados da planilha
for row in sheet.iter_rows(min_row=2, values_only=True):
    nome = row[0]
    cpf = row[1]
    endereco = row[2]

    # Carrega o documento Word
    document = Document('./teste.docx')

    # Cria uma cópia do documento para preencher
    filled_document = copy.deepcopy(document)

    # Preenche os campos no documento Word
    for paragraph in filled_document.paragraphs:
        if '<<NOME>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<NOME>>', nome)
        if '<<CPF>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<CPF>>', str(cpf))
        if '<<ENDERECO>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<ENDERECO>>', endereco)

    # Salva o documento preenchido em um arquivo temporário
    temp_file = tempfile.NamedTemporaryFile(delete=False)
    filled_document.save(temp_file.name)

    # Salva o documento preenchido como PDF
    output_pdf = f'formulario_preenchido_{nome}.pdf'
    convert_to_pdf(temp_file.name, output_pdf)

    # Exclui o arquivo temporário
    os.remove(temp_file.name)